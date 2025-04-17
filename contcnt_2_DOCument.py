'''from docx import Document

def create_docx(description, filename="output.docx"):
    doc = Document()
    doc.add_heading('Generated Description', level=1)
    doc.add_paragraph(description)
    doc.save(filename)
    print(f"DOCX saved as {filename}")'''
import sys
from docx import Document

data=sys.argv[1]
print(f"{data}")


doc=Document()
doc.add_heading("Generated description",level=1)
doc.add_paragraph(data)
doc.save("result.docx")
print("DOCX saved as result.docx")